#!/usr/bin/env python3
"""Inventaría fuentes nuevas sin modificar los archivos originales."""

from __future__ import annotations

import argparse
import os
from collections import Counter, defaultdict
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from PIL import Image
from pypdf import PdfReader


SEARCH_TERMS = (
    "unidad de técnica legislativa",
    "secretario/a de la asamblea",
    "secretario de la asamblea",
    "negocio jurídico solemne",
    "artículo 1453",
    "articulo 1453",
    "cuarto nivel en los ámbitos",
    "cuarto nivel en los ambitos",
)


def normalize(text: str) -> str:
    return " ".join(text.casefold().split())


def docx_text(path: Path) -> tuple[str, int]:
    document = Document(path)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    with ZipFile(path) as archive:
        media_count = sum(name.startswith("word/media/") for name in archive.namelist())
    return "\n".join(chunks), media_count


def matches(text: str) -> list[str]:
    normalized = normalize(text)
    return [term for term in SEARCH_TERMS if term in normalized]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    source = args.source.resolve()
    files = sorted(path for path in source.rglob("*") if path.is_file())
    extension_counts = Counter(path.suffix.casefold() or "(sin extensión)" for path in files)
    subject_counts: dict[str, Counter[str]] = defaultdict(Counter)
    rows: list[dict[str, object]] = []
    findings: list[tuple[str, list[str]]] = []

    for path in files:
        relative = path.relative_to(source)
        subject = relative.parts[0] if len(relative.parts) > 1 else "Archivos en la raíz"
        extension = path.suffix.casefold() or "(sin extensión)"
        subject_counts[subject][extension] += 1
        detail = ""
        found: list[str] = []
        try:
            if extension == ".pdf":
                reader = PdfReader(path)
                texts = [(page.extract_text() or "") for page in reader.pages]
                text = "\n".join(texts)
                image_only_pages = sum(not item.strip() for item in texts)
                detail = f"{len(reader.pages)} pág.; {len(text):,} caracteres; {image_only_pages} pág. sin texto"
                found = matches(text)
            elif extension == ".docx":
                text, media_count = docx_text(path)
                detail = f"{len(text):,} caracteres; {media_count} imágenes incrustadas"
                found = matches(text)
            elif extension in {".jpg", ".jpeg", ".png"}:
                with Image.open(path) as image:
                    detail = f"{image.width}×{image.height} px"
            else:
                detail = f"{path.stat().st_size:,} bytes"
        except Exception as error:  # el reporte debe continuar con archivos dañados
            detail = f"ERROR: {error}"
        rows.append({"path": str(relative), "type": extension, "detail": detail})
        if found:
            findings.append((str(relative), found))

    lines = [
        "# Inventario de fuentes Drive - julio de 2026",
        "",
        f"- Carpeta original: `{source}`",
        f"- Total de archivos: **{len(files)}**",
        "- Los archivos originales no fueron modificados.",
        "",
        "## Resumen por formato",
        "",
    ]
    for extension, count in extension_counts.most_common():
        lines.append(f"- `{extension}`: {count}")
    lines.extend(["", "## Resumen por carpeta", ""])
    for subject in sorted(subject_counts):
        summary = ", ".join(f"{ext}: {count}" for ext, count in subject_counts[subject].most_common())
        lines.append(f"- **{subject}:** {summary}")
    lines.extend(["", "## Coincidencias textuales de las preguntas nuevas", ""])
    if findings:
        for filename, terms in findings:
            lines.append(f"- `{filename}`: {', '.join(terms)}")
    else:
        lines.append("- No se localizaron en las capas de texto de PDF o Word; deben buscarse visualmente en imágenes y páginas renderizadas.")
    lines.extend(["", "## Detalle de archivos", "", "| Archivo | Formato | Detalle |", "|---|---:|---|"])
    for row in rows:
        safe_path = str(row["path"]).replace("|", "\\|")
        safe_detail = str(row["detail"]).replace("|", "\\|")
        lines.append(f"| `{safe_path}` | `{row['type']}` | {safe_detail} |")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(f"{args.output} — {len(files)} archivos")


if __name__ == "__main__":
    main()
